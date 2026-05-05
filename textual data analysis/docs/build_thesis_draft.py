from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "artifacts" / "reports"
OUT = ROOT / "docs" / "智慧財產判決勝敗訴預測_論文初稿.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_east_asia_font(run, font_name="Microsoft JhengHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def add_table(doc, df, caption=None, max_rows=None, percent_cols=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(45, 55, 72)

    if max_rows:
        df = df.head(max_rows)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    header = table.rows[0].cells
    for i, col in enumerate(df.columns):
        set_cell_text(header[i], col, bold=True, color=(255, 255, 255))
        set_cell_shading(header[i], "2F5597")
        header[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            value = row[col]
            if pd.isna(value):
                text = "-"
            elif percent_cols and col in percent_cols:
                text = f"{float(value):.3f}"
            elif isinstance(value, float):
                text = f"{value:.3f}"
            else:
                text = str(value)
            set_cell_text(cells[i], text)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width=5.9):
    image_path = Path(image_path)
    if not image_path.exists():
        p = doc.add_paragraph()
        p.add_run(f"[圖檔待補：{image_path}]").italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(45, 55, 72)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EAF2F8")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(45, 55, 72)
    doc.add_paragraph()


def add_heading_numbered(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return h


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    set_east_asia_font(r)
    return p


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.color.rgb = RGBColor(31, 78, 121)
    return doc


def compute_tables():
    summary = pd.read_csv(REPORTS / "step3_full_runs_summary.csv")
    summary["dataset_label"] = summary["dataset"].map(
        {
            "administrative_win_lose_mixed": "行政",
            "civil_win_lose_mixed": "民事",
            "criminal_win_lose_mixed": "刑事",
            "cwc_win_lose_mixed": "刑事附帶民事",
        }
    )

    valid_summary_path = next(REPORTS.rglob("step3_runs/full_*/valid_target_summary.csv"))
    valid_summary = pd.read_csv(valid_summary_path)
    valid_summary = valid_summary.rename(columns={"JTYPE": "案件類型", "Lose": "Lose", "Mixed": "Mixed", "Win": "Win", "All": "合計"})
    valid_summary["案件類型"] = valid_summary["案件類型"].replace(
        {
            "ADMINISTRATIVE": "行政",
            "CIVIL": "民事",
            "CRIMINAL": "刑事",
            "CWC": "刑事附帶民事",
            "All": "合計",
        }
    )

    proposed = summary.loc[summary.groupby("dataset")["Proposed MNIR + SVM F1"].idxmax()].copy()
    proposed_table = proposed[
        ["dataset_label", "leakage", "feature", "best_chi2_k", "best_svm_c", "Proposed MNIR + SVM Acc", "Proposed MNIR + SVM F1"]
    ].rename(
        columns={
            "dataset_label": "資料集",
            "leakage": "Leakage版本",
            "feature": "特徵",
            "best_chi2_k": "最佳K",
            "best_svm_c": "最佳C",
            "Proposed MNIR + SVM Acc": "Accuracy",
            "Proposed MNIR + SVM F1": "Macro F1",
        }
    )

    long_rows = []
    model_cols = [
        ("Majority Class", "Majority Class Acc", "Majority Class F1"),
        ("BoW + SVM", "BOW + SVM Acc", "BOW + SVM F1"),
        ("TF + SVM", "TF + SVM Acc", "TF + SVM F1"),
        ("TF-IDF + SVM", "TFIDF + SVM Acc", "TFIDF + SVM F1"),
        ("MNIR + SVM", "Proposed MNIR + SVM Acc", "Proposed MNIR + SVM F1"),
    ]
    for _, row in summary.iterrows():
        for model, acc_col, f1_col in model_cols:
            if f1_col in summary.columns and pd.notna(row.get(f1_col)):
                long_rows.append(
                    {
                        "資料集": row["dataset_label"],
                        "Leakage版本": row["leakage"],
                        "特徵": row["feature"],
                        "模型": model,
                        "Accuracy": row[acc_col],
                        "Macro F1": row[f1_col],
                    }
                )
    long_df = pd.DataFrame(long_rows)
    best_overall = long_df.loc[long_df.groupby("資料集")["Macro F1"].idxmax()].sort_values("資料集")

    leakage_compare = (
        long_df.groupby(["資料集", "Leakage版本"])["Macro F1"]
        .max()
        .reset_index()
        .pivot(index="資料集", columns="Leakage版本", values="Macro F1")
        .reset_index()
    )

    return summary, valid_summary, proposed_table.sort_values("資料集"), best_overall, leakage_compare


def find_best_run_images():
    target = (
        REPORTS
        / "cwc_win_lose_mixed"
        / "no_leakage"
        / "bow"
        / "step3_runs"
        / "full_20260504_154912"
    )
    return {
        "heatmap": target / "chi2_k_svm_c_validation_heatmap.png",
        "confusion": target / "final_confusion_matrix.png",
        "comparison": target / "model_comparison_test_metrics.png",
    }


def build():
    summary, valid_summary, proposed_table, best_overall, leakage_compare = compute_tables()
    images = find_best_run_images()
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    r = title.add_run("智慧財產權判決書勝敗訴預測之文字探勘研究")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(31, 78, 121)
    set_east_asia_font(r)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("以卡方特徵選擇、SVM 與 MNIR 特徵轉換為例")
    r.font.size = Pt(15)
    set_east_asia_font(r)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(90)
    meta.add_run("論文初稿\n").bold = True
    meta.add_run("作者：［請填入姓名］\n系所：［請填入系所］\n指導教授：［請填入姓名］\n日期：2026 年 5 月")

    add_note(
        doc,
        "初稿說明：本文件依目前專案程式與 Step 3 full run 結果自動整理，題名、研究問題、文獻引用、學校格式與附錄細節仍需由作者確認後修訂。",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading_numbered(doc, "摘要", level=1)
    add_para(
        doc,
        "本研究以司法院公開之智慧財產權相關判決書為研究對象，建立一套從判決書前處理、中文斷詞、文件詞彙矩陣建構、特徵選擇到模型評估的文字資料分析流程。研究目標為依據判決書中事實段落與相關文本特徵，預測案件結果屬於勝訴、敗訴或部分勝敗訴三類，以評估傳統機器學習方法在法律文本分類任務中的適用性。"
    )
    add_para(
        doc,
        "實驗將案件區分為行政、民事、刑事與刑事附帶民事四類資料集，並比較移除結果洩漏詞彙前後之文本版本。特徵表示包含 BoW、TF 與 TF-IDF，模型流程則比較多數類別基準模型、卡方特徵選擇後之 SVM，以及卡方特徵選擇、MNIR 特徵轉換與 SVM 結合之流程。為避免資料洩漏，所有卡方特徵選擇均僅於訓練集上配適，超參數 K 與 SVM 懲罰參數 C 則依驗證集 Macro F1 選定，測試集僅用於最終評估。"
    )
    add_para(
        doc,
        "實驗結果顯示，資料集存在明顯類別不平衡，僅以 Accuracy 衡量容易高估模型表現，因此本研究以 Macro F1 作為主要指標。整體而言，卡方特徵選擇後直接訓練 SVM 的傳統稀疏文字模型，在四類資料集中皆取得高於 MNIR 流程的 Macro F1，顯示在本研究任務中，經特徵篩選後的詞彙特徵仍保有較強分類訊息。"
    )
    add_para(doc, "關鍵字：智慧財產權、判決預測、文字探勘、卡方特徵選擇、支持向量機、MNIR")

    add_heading_numbered(doc, "目錄（初稿）", level=1)
    for item in [
        "第一章 緒論",
        "第二章 文獻回顧",
        "第三章 研究方法",
        "第四章 實驗結果",
        "第五章 討論",
        "第六章 結論與未來研究",
        "參考文獻與附錄",
    ]:
        doc.add_paragraph(item, style=None)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading_numbered(doc, "第一章 緒論", level=1)
    add_heading_numbered(doc, "1.1 研究背景", level=2)
    add_para(
        doc,
        "司法判決書累積了大量關於法律爭點、事實認定與裁判結果的文本資訊。隨著自然語言處理與機器學習方法逐漸應用於法律資訊分析，如何將非結構化判決文字轉換為可供模型學習的特徵，並建立可重現的分類流程，成為法律科技與實證法學研究的重要議題。智慧財產權案件兼具專業術語密集、案件類型多元與裁判結果高度不平衡等特性，因此適合作為法律文本分類方法的實驗場域。"
    )
    add_heading_numbered(doc, "1.2 研究目的", level=2)
    add_para(
        doc,
        "本研究目的在於建立智慧財產權判決書勝敗訴預測流程，並比較不同文字特徵表示與分類流程的預測表現。具體而言，本研究關注三個問題：第一，經過規則標註與中文斷詞後，傳統詞彙特徵是否能有效預測判決結果；第二，卡方特徵選擇與 SVM 是否足以形成具競爭力之基準模型；第三，MNIR 特徵轉換是否能進一步改善模型表現。"
    )
    add_heading_numbered(doc, "1.3 研究範圍與限制", level=2)
    add_para(
        doc,
        "本研究以已蒐集並前處理之智慧財產權相關判決資料為範圍，標籤聚焦於 Win、Lose 與 Mixed 三類。研究不主張模型可取代法律專業判斷，而是將其定位為文本分類與實證分析工具。由於標籤係由規則與人工例外修正建立，仍可能存在標註誤差；此外，案件文字中的裁判結果詞彙可能形成 outcome leakage，因此本研究另行建立 no-leakage 版本以檢驗其影響。"
    )

    add_heading_numbered(doc, "第二章 文獻回顧", level=1)
    add_para(
        doc,
        "本章待補正式文獻引用。初稿建議分為三類文獻：法律文本探勘與裁判預測、中文法律文本斷詞與特徵表示、以及文字分類模型與特徵選擇方法。最終版本應補入國內外法律判決預測、SVM 於文本分類之應用、TF-IDF 與卡方特徵選擇，以及 MNIR 或相關監督式特徵轉換方法之研究。"
    )
    add_note(doc, "待確認問題：是否已有指定文獻、學校要求引用格式，或需納入法學方法論相關文獻。")

    add_heading_numbered(doc, "第三章 研究方法", level=1)
    add_heading_numbered(doc, "3.1 研究流程", level=2)
    add_para(
        doc,
        "本研究流程包含資料蒐集、資料前處理、文本特徵建構、模型訓練與模型評估五個階段。首先以智慧財產相關關鍵字與案件資訊篩選判決書；其次透過規則判定案件類型與裁判結果，並自全文擷取事實段落；接著使用 CKIP 進行中文斷詞，並建立 BoW、TF 與 TF-IDF 三種文件詞彙矩陣；最後以分層資料切分、卡方特徵選擇、SVM 與 MNIR 流程進行實驗比較。"
    )
    add_heading_numbered(doc, "3.2 資料標籤與案件類型", level=2)
    add_para(
        doc,
        "本研究將有效案件依 JTYPE 區分為行政、民事、刑事與刑事附帶民事四類，並將裁判結果整理為 Win、Lose 與 Mixed。資料前處理階段會排除裁定、不重要案件與非目標結果，以形成可供監督式學習使用的目標資料集。"
    )
    add_table(doc, valid_summary, "表 1 目標資料集之案件類型與裁判結果分布")

    add_heading_numbered(doc, "3.3 文本特徵表示", level=2)
    add_para(
        doc,
        "本研究比較三種常見文字特徵表示。BoW 以詞彙出現次數表示文件；TF 將詞頻正規化，以降低文件長度差異造成的影響；TF-IDF 則進一步考量詞彙在語料庫中的辨識程度。三種特徵皆以相同資料切分與超參數選擇流程進行比較。"
    )
    add_heading_numbered(doc, "3.4 卡方特徵選擇", level=2)
    add_para(
        doc,
        "由於法律文本 DTM 維度高且包含大量低資訊詞彙，本研究採用卡方統計量篩選與分類標籤較相關的詞彙特徵。為避免資料洩漏，卡方選擇僅於訓練集上配適，再套用至驗證集與測試集。特徵數 K 不預先固定，而是以驗證集 Macro F1 於候選值中選定。"
    )
    add_heading_numbered(doc, "3.5 模型與基準方法", level=2)
    add_para(
        doc,
        "本研究設計三類比較模型。第一為 Majority Class baseline，永遠預測訓練集中最多的類別，用以衡量資料不平衡下的最低基準。第二為卡方特徵選擇後直接訓練 SVM 的稀疏文字模型。第三為本研究原先設定之 MNIR-based pipeline，即先以 MNIR 轉換特徵後，再以 SVM 進行分類。"
    )
    add_heading_numbered(doc, "3.6 評估設計", level=2)
    add_para(
        doc,
        "資料採 70%、10%、20% 之分層訓練集、驗證集與測試集切分。卡方 K 與 SVM C 僅依驗證集 Macro F1 選定，測試集僅在最終模型確定後評估一次。評估指標包含 Accuracy、Macro F1、各類別 Precision、Recall、F1-score，以及混淆矩陣。考量本研究資料類別分布不均，後續結果以 Macro F1 作為主要判斷依據。"
    )

    add_heading_numbered(doc, "第四章 實驗結果", level=1)
    add_heading_numbered(doc, "4.1 資料不平衡現象", level=2)
    add_para(
        doc,
        "由表 1 可見，各案件類型均呈現明顯類別不平衡。行政與刑事案件中 Lose 類別占比最高；民事案件雖有較多 Mixed 案件，但 Win 仍為少數；刑事附帶民事案件則以 Mixed 與 Lose 為主，Win 樣本極少。此現象使 Accuracy 容易受多數類別主導，因此 Macro F1 更能反映模型於少數類別上的分類能力。"
    )
    add_heading_numbered(doc, "4.2 MNIR + SVM 流程之最佳結果", level=2)
    add_table(doc, proposed_table, "表 2 MNIR + SVM 流程於各資料集之最佳測試結果", percent_cols=["Accuracy", "Macro F1"])
    add_para(
        doc,
        "MNIR + SVM 流程在四類資料集中的最佳 Macro F1 介於 0.318 至 0.571。刑事附帶民事資料集表現相對較佳，行政、民事與刑事資料集則因類別不平衡與少數類別樣本不足，Macro F1 顯著低於 Accuracy。"
    )
    add_heading_numbered(doc, "4.3 Baseline 與模型比較", level=2)
    best_overall_display = best_overall[["資料集", "Leakage版本", "特徵", "模型", "Accuracy", "Macro F1"]]
    add_table(doc, best_overall_display, "表 3 各資料集最佳模型比較", percent_cols=["Accuracy", "Macro F1"])
    add_para(
        doc,
        "實驗結果顯示，四類資料集之最佳 Macro F1 均由卡方特徵選擇後直接訓練 SVM 的模型取得。此結果表示，在本研究資料中，傳統稀疏詞彙特徵經適當特徵篩選後，能保留較多與判決結果相關的分類訊息；相較之下，MNIR 轉換後的特徵並未穩定改善分類效果。"
    )
    add_table(doc, leakage_compare, "表 4 各資料集 with-leakage 與 no-leakage 版本之最佳 Macro F1", percent_cols=["no_leakage", "with_leakage"])

    add_heading_numbered(doc, "4.4 超參數敏感度與參數高原", level=2)
    add_para(
        doc,
        "本研究以 validation Macro F1 選擇卡方特徵數 K 與 SVM 參數 C，並保留完整 K × C 參數表面。參數熱圖可用於判斷模型表現是否集中於單一最佳點，或在相鄰參數區間呈現穩定高原。若高分區域較寬，代表模型對該超參數設定較不敏感；反之，若高分僅出現在單一格點，則需更謹慎解讀模型穩定性。"
    )
    add_figure(doc, images["heatmap"], "圖 1 參數表面範例：Chi-square K 與 SVM C 之 Validation Macro F1", width=5.8)
    add_figure(doc, images["comparison"], "圖 2 模型比較範例：測試集 Accuracy 與 Macro F1", width=5.8)
    doc.add_page_break()
    add_heading_numbered(doc, "4.5 混淆矩陣分析", level=2)
    add_para(
        doc,
        "混淆矩陣可進一步檢視模型錯誤型態。由於多數資料集中 Lose 類別樣本較多，模型可能傾向將少數類別 Win 或 Mixed 誤判為 Lose。此現象說明在法律判決預測任務中，僅改善整體準確率不足以代表模型具有均衡分類能力，後續研究應進一步考慮類別權重、重抽樣或代價敏感學習。"
    )
    add_figure(doc, images["confusion"], "圖 3 最佳模型範例之測試集混淆矩陣", width=5.3)

    add_heading_numbered(doc, "第五章 討論", level=1)
    add_heading_numbered(doc, "5.1 主要發現", level=2)
    add_para(
        doc,
        "本研究主要發現有三點。第一，智慧財產權判決資料存在嚴重類別不平衡，Accuracy 高低不能單獨代表分類品質。第二，卡方特徵選擇與 SVM 組成的傳統文字分類流程在本資料上具備強競爭力，且在四類資料集中均優於 MNIR + SVM 流程。第三，移除 outcome leakage 後，部分資料集表現下降但並非一致，顯示結果洩漏詞彙對不同案件類型與特徵表示的影響程度有所差異。"
    )
    add_heading_numbered(doc, "5.2 MNIR 未優於 Baseline 之可能原因", level=2)
    add_para(
        doc,
        "MNIR 特徵轉換未能優於直接使用卡方篩選後稀疏特徵，可能有數個原因。首先，法律判決結果與特定詞彙或片語之關聯可能已能由高維稀疏特徵直接捕捉，轉換後反而損失細部訊息。其次，少數類別樣本過少，使監督式特徵轉換難以穩定學習類別差異。最後，不同案件類型的裁判語言結構差異大，單一轉換流程未必適合所有資料集。"
    )
    add_heading_numbered(doc, "5.3 實務與研究意涵", level=2)
    add_para(
        doc,
        "本研究結果提醒，在法律文本機器學習任務中，較複雜的特徵轉換方法不必然優於簡潔且可解釋的傳統模型。對於研究者而言，建立完整 baseline、保留參數選擇紀錄並以 Macro F1 與混淆矩陣檢驗模型，是避免過度解讀模型表現的重要步驟。對實務應用而言，模型輸出應被視為案件文本分析輔助，而非對個案結果之法律預測或判斷。"
    )
    add_heading_numbered(doc, "5.4 研究限制", level=2)
    add_para(
        doc,
        "本研究仍有數項限制。第一，裁判結果標籤由規則與人工例外修正產生，仍需進一步抽樣驗證標籤正確性。第二，本文尚未納入深度語言模型或法律專用語言模型，因此無法比較傳統模型與預訓練模型之差異。第三，本研究以三類結果標籤為主，尚未細緻處理和解、發回、更審或程序性駁回等法律結果。第四，各案件類型之樣本數與類別比例差異大，模型比較仍受資料結構影響。"
    )

    add_heading_numbered(doc, "第六章 結論與未來研究", level=1)
    add_para(
        doc,
        "本研究建立一套智慧財產權判決書文字分析與勝敗訴分類流程，涵蓋判決資料前處理、中文斷詞、DTM 特徵建構、卡方特徵選擇、SVM 分類、MNIR 特徵轉換與完整模型評估。實驗結果顯示，在目前資料與設定下，卡方特徵選擇後之 SVM baseline 整體表現優於 MNIR + SVM，尤其在 Macro F1 指標上更能反映少數類別分類能力。"
    )
    add_para(
        doc,
        "未來研究可從四個方向延伸。第一，進行更嚴謹的標籤抽樣檢查與錯誤分析。第二，加入 class weighting、重抽樣或 focal loss 等方法處理類別不平衡。第三，嘗試法律語言模型或中文預訓練模型，比較其與傳統稀疏特徵模型的差異。第四，將模型解釋方法納入分析，例如檢視高卡方詞彙、SVM 權重或 SHAP 值，以提升法律文本分類結果的可解釋性。"
    )

    add_heading_numbered(doc, "參考文獻（待補）", level=1)
    for ref in [
        "［待補］法律文本探勘與裁判預測相關文獻。",
        "［待補］SVM 與傳統文字分類方法相關文獻。",
        "［待補］卡方特徵選擇、TF-IDF 與中文斷詞相關文獻。",
        "［待補］MNIR 或監督式文字特徵轉換相關文獻。",
    ]:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)

    add_heading_numbered(doc, "附錄：待確認事項", level=1)
    for item in [
        "正式論文題名、作者姓名、學校、系所與指導教授資訊。",
        "學校規定之頁面格式、摘要格式、章節編號與參考文獻格式。",
        "是否需要加入英文摘要與英文關鍵字。",
        "是否要以 BoW + SVM 作為主模型，將 MNIR + SVM 放入消融比較。",
        "是否需要補入完整四類資料集之混淆矩陣與參數熱圖作為附錄。",
        "文獻回顧需補入之指定文獻清單。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
