from pathlib import Path

import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from build_thesis_draft import (
    REPORTS,
    ROOT,
    add_figure,
    add_heading_numbered,
    add_note,
    add_para,
    add_table,
    setup_document,
    set_east_asia_font,
)


OUT = ROOT / "docs" / "論文第三章以後_修正版.docx"
RUN_TAG = "full_20260504_154912"

DATASET_LABELS = {
    "administrative_win_lose_mixed": "行政",
    "civil_win_lose_mixed": "民事",
    "criminal_win_lose_mixed": "刑事",
    "cwc_win_lose_mixed": "刑事附帶民事",
}

FEATURE_LABELS = {"bow": "BoW", "tf": "TF", "tfidf": "TF-IDF"}
DIRECT_MODELS = {"BOW + SVM", "TF + SVM", "TFIDF + SVM"}


def normalize_model(name):
    return name.replace("BOW", "BoW").replace("TFIDF", "TF-IDF").replace("Proposed ", "")


def run_dir(dataset, leakage, feature):
    return REPORTS / dataset / leakage / feature / "step3_runs" / RUN_TAG


def load_results():
    rows = []
    for path in REPORTS.glob(f"*/*/*/step3_runs/{RUN_TAG}/model_comparison.csv"):
        feature = path.parents[2].name
        leakage = path.parents[3].name
        dataset = path.parents[4].name
        cfg_path = path.parent / "run_config.csv"
        cfg = pd.read_csv(cfg_path).iloc[0] if cfg_path.exists() else {}
        df = pd.read_csv(path)
        for _, row in df.iterrows():
            model = str(row["Model"])
            model_norm = normalize_model(model)
            is_direct = model in DIRECT_MODELS
            rows.append(
                {
                    "dataset": dataset,
                    "資料集": DATASET_LABELS.get(dataset, dataset),
                    "leakage": leakage,
                    "Leakage": "no leakage" if leakage == "no_leakage" else "with leakage",
                    "feature": feature,
                    "特徵": FEATURE_LABELS.get(feature, feature),
                    "模型": model_norm,
                    "模型群組": "Direct SVM" if is_direct else model_norm,
                    "Accuracy": row["Test Accuracy"],
                    "Macro F1": row["Test Macro F1"],
                    "Best C": row.get("Best C"),
                    "Best K": cfg.get("best_direct_chi2_k") if is_direct else cfg.get("best_chi2_k"),
                    "run_dir": path.parent,
                }
            )
    long_df = pd.DataFrame(rows)

    target_summary_path = next(REPORTS.rglob(f"step3_runs/{RUN_TAG}/valid_target_summary.csv"))
    target_summary = pd.read_csv(target_summary_path).rename(columns={"JTYPE": "資料集", "All": "合計"})
    target_summary["資料集"] = target_summary["資料集"].replace(
        {
            "ADMINISTRATIVE": "行政",
            "CIVIL": "民事",
            "CRIMINAL": "刑事",
            "CWC": "刑事附帶民事",
            "All": "合計",
        }
    )

    direct = long_df[long_df["模型群組"] == "Direct SVM"].copy()
    direct_best = direct.loc[direct.groupby("資料集")["Macro F1"].idxmax()].sort_values("資料集")

    feature_compare = (
        direct.groupby(["資料集", "特徵"])["Macro F1"]
        .max()
        .reset_index()
        .pivot(index="資料集", columns="特徵", values="Macro F1")
        .reset_index()
    )
    feature_compare = feature_compare[["資料集", "BoW", "TF", "TF-IDF"]]

    mnir_best = (
        long_df[long_df["模型"] == "MNIR + SVM"]
        .loc[lambda df: df.groupby("資料集")["Macro F1"].idxmax()]
        .sort_values("資料集")
    )
    majority_best = (
        long_df[long_df["模型"] == "Majority Class"]
        .loc[lambda df: df.groupby("資料集")["Macro F1"].idxmax()]
        .sort_values("資料集")
    )

    overall = direct_best[["資料集", "Leakage", "特徵", "Accuracy", "Macro F1"]].rename(
        columns={"Leakage": "Direct Leakage", "特徵": "Direct 特徵"}
    )
    overall = overall.merge(
        mnir_best[["資料集", "Macro F1"]].rename(columns={"Macro F1": "MNIR+SVM Macro F1"}),
        on="資料集",
        how="left",
    )
    overall = overall.merge(
        majority_best[["資料集", "Macro F1"]].rename(columns={"Macro F1": "Majority Macro F1"}),
        on="資料集",
        how="left",
    )
    overall["Direct - MNIR"] = overall["Macro F1"] - overall["MNIR+SVM Macro F1"]

    baseline_compare = direct_best[["資料集", "Leakage", "特徵", "模型", "Macro F1"]].rename(
        columns={"Leakage": "Direct最佳Leakage", "特徵": "Direct最佳特徵", "模型": "Direct模型", "Macro F1": "Direct SVM Macro F1"}
    )
    baseline_compare = baseline_compare.merge(
        mnir_best[["資料集", "Leakage", "特徵", "Macro F1"]].rename(
            columns={"Leakage": "MNIR最佳Leakage", "特徵": "MNIR最佳特徵", "Macro F1": "MNIR+SVM Macro F1"}
        ),
        on="資料集",
        how="left",
    )
    baseline_compare["差異"] = baseline_compare["Direct SVM Macro F1"] - baseline_compare["MNIR+SVM Macro F1"]

    leakage_compare = (
        direct.groupby(["資料集", "Leakage"])["Macro F1"]
        .max()
        .reset_index()
        .pivot(index="資料集", columns="Leakage", values="Macro F1")
        .reset_index()
    )
    leakage_compare = leakage_compare[["資料集", "no leakage", "with leakage"]]
    leakage_compare["no - with"] = leakage_compare["no leakage"] - leakage_compare["with leakage"]

    best_row = direct.loc[direct["Macro F1"].idxmax()]
    best_dir = Path(best_row["run_dir"])
    k_summary = pd.read_csv(best_dir / "direct_svm_chi2_k_tuning_summary.csv")
    best_val = k_summary["Validation Macro F1"].max()
    k_summary["與最佳差距"] = best_val - k_summary["Validation Macro F1"]
    plateau = k_summary[k_summary["與最佳差距"] <= 0.01].copy()
    plateau_table = k_summary[["K", "Selected Features", "Best C", "Validation Macro F1", "與最佳差距"]].head(6).copy()
    plateau_table["K"] = plateau_table["K"].astype(int).astype(str)
    plateau_table["Selected Features"] = plateau_table["Selected Features"].astype(int).astype(str)

    feature_slug = best_row["feature"]
    cm = pd.read_csv(best_dir / f"baseline_{feature_slug}_svm_confusion_matrix.csv", index_col=0).reset_index()
    cm = cm.rename(columns={cm.columns[0]: "實際\\預測"})
    report = pd.read_csv(best_dir / f"baseline_{feature_slug}_svm_test_report.csv").rename(columns={"Unnamed: 0": "類別"})
    report = report[report["類別"].isin(["Lose", "Mixed", "Win", "macro avg", "weighted avg"])]

    return {
        "long": long_df,
        "target_summary": target_summary,
        "direct_best": direct_best,
        "feature_compare": feature_compare,
        "overall": overall,
        "baseline_compare": baseline_compare,
        "leakage_compare": leakage_compare,
        "best_row": best_row,
        "best_dir": best_dir,
        "plateau_table": plateau_table,
        "plateau_count": len(plateau),
        "confusion_matrix": cm,
        "class_report": report,
    }


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("第三章以後修正版")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)
    set_east_asia_font(r)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run("以 BoW、TF、TF-IDF + SVM 比較為主軸，MNIR+SVM 作為 baseline")
    rr.font.size = Pt(12)
    set_east_asia_font(rr)


def build():
    data = load_results()
    doc = setup_document()
    add_title(doc)
    add_note(doc, "本文件只包含第三章以後內容，可接續於既有緒論與文獻回顧之後。第四章已改為 Empirical Results 的研究問題導向結構。")

    add_heading_numbered(doc, "第三章 研究方法", level=1)
    add_heading_numbered(doc, "3.1 研究流程", level=2)
    add_para(
        doc,
        "本研究以智慧財產權相關判決書為研究對象，將裁判結果整理為 Lose、Mixed 與 Win 三類標籤，並依案件類型切分為行政、民事、刑事與刑事附帶民事四個資料集。整體流程包含資料前處理、斷詞與文件詞項矩陣建置、卡方特徵選擇、模型訓練、validation 參數選擇與 test set 評估。"
    )
    add_para(
        doc,
        "實驗設計的核心不是提出單一新模型，而是比較三種傳統稀疏文字表示法 BoW、TF 與 TF-IDF 在法律判決勝敗訴預測上的表現。Direct SVM 為主要比較流程，MNIR+SVM 則作為監督式特徵轉換 baseline，用以檢驗 MNIR 是否能改善傳統稀疏特徵。"
    )

    add_heading_numbered(doc, "3.2 資料與標籤分布", level=2)
    add_para(
        doc,
        "資料標籤呈現明顯不平衡，部分案件類型中 Lose 或 Mixed 為主要類別，而 Win 類別樣本數較少。因此，本文除報告 Accuracy 外，也以 Macro F1 作為主要評估指標，以避免模型僅因預測多數類別而獲得表面上較高的準確率。"
    )
    add_table(doc, data["target_summary"], "表 1 資料集與裁判結果標籤分布")

    add_heading_numbered(doc, "3.3 特徵表示與卡方選擇", level=2)
    add_para(
        doc,
        "本文比較 BoW、TF 與 TF-IDF 三種文字表示。所有表示法皆先產生文件詞項矩陣，再於訓練資料上進行卡方特徵選擇；validation 與 test 資料只套用訓練階段選出的詞項，避免評估資料洩漏至特徵選擇過程。"
    )
    add_para(
        doc,
        "卡方特徵數 K 不採固定值，而是在 validation set 上與 SVM 之 C 共同選擇。此設定可避免任意指定 K 造成結論偏誤，也讓不同資料集與不同表示法能選擇各自較合適的特徵空間大小。"
    )

    add_heading_numbered(doc, "3.4 模型與 Baseline", level=2)
    add_para(
        doc,
        "Direct SVM 流程為 BoW、TF 或 TF-IDF 經卡方特徵選擇後直接訓練線性 SVM。MNIR+SVM baseline 則在同樣的卡方選擇後進行 MNIR 特徵轉換，再訓練線性 SVM。另加入 Majority Class 作為最低基準，用以確認模型是否真正超越多數類別預測。"
    )

    add_heading_numbered(doc, "3.5 評估設定", level=2)
    add_para(
        doc,
        "資料依 stratified split 分為 training、validation 與 test sets。參數 K 與 C 皆以 validation Macro F1 選擇，最後僅在 test set 上報告模型泛化表現。主要指標為 Macro F1，輔以 Accuracy、各類別 precision、recall、F1-score 與混淆矩陣。"
    )

    add_heading_numbered(doc, "第四章 Empirical Results", level=1)
    add_heading_numbered(doc, "4.1 Descriptive Statistics of Dataset", level=2)
    add_para(
        doc,
        "標籤分布顯示本任務具有類別不平衡特性。若僅觀察 Accuracy，模型可能因偏向 Lose 或 Mixed 等多數類別而看似表現良好；因此後續結果以 Macro F1 為主，並在 error analysis 中檢查少數類別是否被系統性忽略。"
    )

    add_heading_numbered(doc, "4.2 Overall Model Performance", level=2)
    add_table(
        doc,
        data["overall"],
        "表 2 各資料集 Direct SVM 最佳結果與 Baseline 比較",
        percent_cols=["Accuracy", "Macro F1", "MNIR+SVM Macro F1", "Majority Macro F1", "Direct - MNIR"],
    )
    add_para(
        doc,
        "整體結果顯示，Direct SVM 在四個資料集皆高於 MNIR+SVM 與 Majority Class。這表示在目前資料與標籤設定下，經卡方篩選後的原始稀疏詞彙特徵已能保留相當多與裁判結果相關的分類訊息。"
    )

    add_heading_numbered(doc, "4.3 Feature Representation Comparison", level=2)
    add_table(doc, data["feature_compare"], "表 3 Direct SVM 下 BoW、TF、TF-IDF 之最佳 Macro F1", percent_cols=["BoW", "TF", "TF-IDF"])
    add_para(
        doc,
        "三種表示法沒有單一方法在所有案件類型皆勝出。刑事資料集以 TF 表現較佳，民事資料集以 TF-IDF 較佳，行政與刑事附帶民事則以 BoW 較佳。此結果說明法律判決文本的有效特徵形式會受案件類型與詞彙分布影響，不能假設 TF-IDF 或任一表示法必然最佳。"
    )

    add_heading_numbered(doc, "4.4 Baseline Comparison: Direct SVM vs. MNIR+SVM", level=2)
    add_table(
        doc,
        data["baseline_compare"],
        "表 4 Direct SVM 與 MNIR+SVM 於各資料集之最佳 Macro F1 比較",
        percent_cols=["Direct SVM Macro F1", "MNIR+SVM Macro F1", "差異"],
    )
    add_para(
        doc,
        "MNIR+SVM 原本可被視為一種監督式特徵轉換 baseline，理論上可能透過類別資訊重新加權詞彙特徵。然而實驗結果未支持此假設。可能原因是判決文本中的判別性詞彙已可由卡方選擇直接捕捉，MNIR 轉換反而壓縮或平滑少數類別訊號，導致 Macro F1 下降。"
    )

    add_heading_numbered(doc, "4.5 Parameter Selection and Validation Plateau", level=2)
    add_para(
        doc,
        f"以最佳整體設定為例，validation 參數表面顯示最佳組合附近存在多個接近最佳的 K 與 C 組合；在與最佳 Validation Macro F1 差距不超過 0.01 的條件下，共有 {data['plateau_count']} 組參數。這表示結果並非完全依賴單一任意 K 值，而是存在相對穩定的參數區域。"
    )
    add_table(
        doc,
        data["plateau_table"],
        "表 5 最佳整體設定之 Direct SVM 參數選擇摘要",
        percent_cols=["Validation Macro F1", "與最佳差距"],
    )
    add_figure(
        doc,
        data["best_dir"] / "direct_svm_chi2_k_svm_c_validation_heatmap.png",
        "圖 1 Direct SVM K × C Validation Macro F1 參數表面",
        width=5.8,
    )

    add_heading_numbered(doc, "4.6 Leakage Analysis", level=2)
    add_table(
        doc,
        data["leakage_compare"],
        "表 6 Direct SVM 下 no-leakage 與 with-leakage 之最佳 Macro F1",
        percent_cols=["no leakage", "with leakage", "no - with"],
    )
    add_para(
        doc,
        "Leakage 移除的影響並不一致。部分資料集在 no-leakage 版本仍維持或取得較佳表現，表示模型並非完全依賴明顯裁判結果詞彙；但行政資料集中 with-leakage 表現較高，顯示 outcome leakage 仍可能影響模型評估。因此，本文主要結論應優先依據 no-leakage 結果解讀。"
    )

    add_heading_numbered(doc, "4.7 Error Analysis", level=2)
    add_para(
        doc,
        "最佳整體模型雖然取得較高 Accuracy 與 Macro F1，但混淆矩陣與類別報告顯示 Win 類別仍是主要困難來源。此現象與資料不平衡一致，表示模型對少數類別的召回能力不足，未來可進一步加入 class weighting、resampling 或針對少數類別的錯誤案例分析。"
    )
    add_table(doc, data["confusion_matrix"], "表 7 最佳整體 Direct SVM 模型之測試集混淆矩陣")
    add_table(
        doc,
        data["class_report"],
        "表 8 最佳整體 Direct SVM 模型之分類報告",
        percent_cols=["precision", "recall", "f1-score"],
    )

    add_heading_numbered(doc, "第五章 討論", level=1)
    add_heading_numbered(doc, "5.1 主要發現", level=2)
    add_para(
        doc,
        "本研究的主要發現有三點。第一，Direct SVM 在各資料集皆優於 MNIR+SVM baseline。第二，BoW、TF 與 TF-IDF 的最佳表現依案件類型而不同。第三，Accuracy 受多數類別影響明顯，Macro F1 與混淆矩陣更能反映模型對少數類別的限制。"
    )

    add_heading_numbered(doc, "5.2 MNIR+SVM 作為 Baseline 的意義", level=2)
    add_para(
        doc,
        "MNIR+SVM 在本文中應定位為比較基準，而非主方法。其結果有助於說明監督式特徵轉換在本資料設定下未必優於傳統稀疏詞彙表示。這個結果本身具有研究意義，因為它提醒法律文本分類不一定需要更複雜的特徵轉換流程，嚴謹調參的簡潔 baseline 可能已具有相當競爭力。"
    )

    add_heading_numbered(doc, "5.3 研究限制", level=2)
    add_para(
        doc,
        "本研究仍受資料標籤品質、類別不平衡、缺乏外部驗證資料與尚未納入大型語言模型等限制。尤其 Win 類別樣本不足，使模型難以學習穩定決策邊界。後續研究可進一步檢查標籤正確性、擴增少數類別樣本，並比較法律語言模型與傳統稀疏特徵模型之差異。"
    )

    add_heading_numbered(doc, "第六章 結論", level=1)
    add_para(
        doc,
        "本文比較 BoW、TF、TF-IDF 三種文字表示與 Direct SVM、MNIR+SVM 兩種流程於智慧財產權判決勝敗訴預測之表現。實驗結果顯示，Direct SVM 在本研究設定下整體優於 MNIR+SVM，且最佳文字表示法會隨案件類型而改變。此發現支持以嚴謹調參的傳統稀疏特徵模型作為法律文本分類研究的強 baseline，並指出未來若導入更複雜模型，仍應與此類 baseline 進行公平比較。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
