from pathlib import Path

import pandas as pd
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from build_thesis_draft import (
    REPORTS,
    ROOT,
    add_figure,
    add_heading_numbered,
    add_note,
    add_para,
    add_table,
    setup_document,
    set_east_asia_font,
)


OUT = ROOT / "docs" / "智慧財產判決勝敗訴預測_論文初稿_v2.docx"


DATASET_LABELS = {
    "administrative_win_lose_mixed": "行政",
    "civil_win_lose_mixed": "民事",
    "criminal_win_lose_mixed": "刑事",
    "cwc_win_lose_mixed": "刑事附帶民事",
}


def load_results():
    pivot_path = REPORTS / "step4_patched_model_comparison_summary.csv"
    if not pivot_path.exists():
        raise FileNotFoundError(f"Missing patched summary: {pivot_path}")

    pivot = pd.read_csv(pivot_path)
    pivot["資料集"] = pivot["dataset"].map(DATASET_LABELS)
    pivot["Leakage版本"] = pivot["leakage"].replace({"with_leakage": "with leakage", "no_leakage": "no leakage"})
    pivot["特徵"] = pivot["feature"].replace({"bow": "BoW", "tf": "TF", "tfidf": "TF-IDF"})

    long_rows = []
    for _, row in pivot.iterrows():
        for model in ["BOW + SVM", "TF + SVM", "TFIDF + SVM", "Proposed MNIR + SVM", "Majority Class"]:
            if model in pivot.columns and pd.notna(row.get(model)):
                long_rows.append(
                    {
                        "dataset": row["dataset"],
                        "資料集": row["資料集"],
                        "Leakage版本": row["Leakage版本"],
                        "特徵": row["特徵"],
                        "模型": model.replace("BOW", "BoW").replace("TFIDF", "TF-IDF").replace("Proposed ", ""),
                        "Macro F1": row[model],
                    }
                )
    long_df = pd.DataFrame(long_rows)

    valid_summary_path = next(REPORTS.rglob("step3_runs/full_20260504_154912/valid_target_summary.csv"))
    valid_summary = pd.read_csv(valid_summary_path)
    valid_summary = valid_summary.rename(columns={"JTYPE": "案件類型", "All": "合計"})
    valid_summary["案件類型"] = valid_summary["案件類型"].replace(
        {
            "ADMINISTRATIVE": "行政",
            "CIVIL": "民事",
            "CRIMINAL": "刑事",
            "CWC": "刑事附帶民事",
            "All": "合計",
        }
    )

    direct = long_df[long_df["模型"].isin(["BoW + SVM", "TF + SVM", "TF-IDF + SVM"])].copy()
    direct_best = direct.loc[direct.groupby("資料集")["Macro F1"].idxmax()].sort_values("資料集")

    feature_compare = (
        direct.groupby(["資料集", "特徵"])["Macro F1"]
        .max()
        .reset_index()
        .pivot(index="資料集", columns="特徵", values="Macro F1")
        .reset_index()
    )
    feature_compare = feature_compare[["資料集", "BoW", "TF", "TF-IDF"]]

    mnir_best = (
        long_df[long_df["模型"] == "MNIR + SVM"]
        .loc[lambda df: df.groupby("資料集")["Macro F1"].idxmax()]
        .sort_values("資料集")
    )
    direct_vs_mnir = direct_best[["資料集", "Leakage版本", "特徵", "模型", "Macro F1"]].rename(
        columns={"Leakage版本": "Direct最佳Leakage", "特徵": "Direct最佳特徵", "模型": "Direct模型", "Macro F1": "Direct SVM Macro F1"}
    )
    mnir_small = mnir_best[["資料集", "Leakage版本", "特徵", "Macro F1"]].rename(
        columns={"Leakage版本": "MNIR最佳Leakage", "特徵": "MNIR最佳特徵", "Macro F1": "MNIR+SVM Macro F1"}
    )
    direct_vs_mnir = direct_vs_mnir.merge(mnir_small, on="資料集", how="left")
    direct_vs_mnir["差異"] = direct_vs_mnir["Direct SVM Macro F1"] - direct_vs_mnir["MNIR+SVM Macro F1"]

    leakage_compare = (
        direct.groupby(["資料集", "Leakage版本"])["Macro F1"]
        .max()
        .reset_index()
        .pivot(index="資料集", columns="Leakage版本", values="Macro F1")
        .reset_index()
    )

    best_row = direct.loc[direct["Macro F1"].idxmax()]
    best_run_dir = (
        REPORTS
        / best_row["dataset"]
        / ("no_leakage" if best_row["Leakage版本"] == "no leakage" else "with_leakage")
        / best_row["特徵"].lower().replace("-", "")
        / "step3_runs"
        / "full_20260504_154912"
    )

    return {
        "pivot": pivot,
        "long": long_df,
        "valid_summary": valid_summary,
        "direct_best": direct_best,
        "feature_compare": feature_compare,
        "direct_vs_mnir": direct_vs_mnir,
        "leakage_compare": leakage_compare,
        "best_row": best_row,
        "best_run_dir": best_run_dir,
    }


def add_best_confusion_matrix(doc, data):
    feature_slug = {"BoW": "bow", "TF": "tf", "TF-IDF": "tfidf"}[data["best_row"]["特徵"]]
    cm_path = data["best_run_dir"] / f"baseline_{feature_slug}_svm_confusion_matrix.csv"
    if not cm_path.exists():
        add_note(doc, f"Direct SVM 混淆矩陣待補：尚未找到 {cm_path.name}。")
        return

    cm = pd.read_csv(cm_path, index_col=0).reset_index()
    cm = cm.rename(columns={cm.columns[0]: "實際\\預測"})
    doc.add_page_break()
    add_table(doc, cm, "表 5 Direct SVM 最佳設定之測試集混淆矩陣")


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    run = title.add_run("智慧財產權判決書勝敗訴預測之文字特徵比較研究")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_east_asia_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("以 BoW、TF、TF-IDF 與 SVM 分類流程為核心")
    run.font.size = Pt(15)
    set_east_asia_font(run)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(90)
    meta.add_run("論文初稿 v2\n").bold = True
    meta.add_run("作者：［請填入姓名］\n系所：［請填入系所］\n指導教授：［請填入姓名］\n日期：2026 年 5 月")

    add_note(
        doc,
        "初稿說明：本版依 Step 4 修正後結果重設論文主軸。Direct SVM 已獨立進行卡方 K 與 SVM C 調參；MNIR+SVM 作為比較基準，而非主模型。",
    )


def build():
    data = load_results()
    doc = setup_document()
    add_cover(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading_numbered(doc, "摘要", level=1)
    add_para(
        doc,
        "本研究以智慧財產權相關判決書為研究對象，建立從文本前處理、中文斷詞、文件詞彙矩陣建構、特徵選擇到分類評估之完整文字探勘流程。研究目的在於比較 BoW、TF 與 TF-IDF 三種文字特徵表示，在智慧財產判決結果分類任務中的表現，並檢驗傳統稀疏文字特徵結合 SVM 是否能在不平衡法律文本資料中取得穩定分類效果。"
    )
    add_para(
        doc,
        "實驗將案件區分為行政、民事、刑事與刑事附帶民事四類，並針對 with-leakage 與 no-leakage 兩種文本版本進行比較。模型流程包含 Majority Class 最低基準、BoW/TF/TF-IDF 經卡方特徵選擇後直接訓練 SVM，以及 MNIR+SVM 基準流程。所有模型均採分層訓練、驗證與測試集切分；卡方特徵數 K 與 SVM 參數 C 均依驗證集 Macro F1 獨立選定，測試集僅用於最終評估。"
    )
    add_para(
        doc,
        "實驗結果顯示，Direct SVM 流程在四類資料集中均優於 MNIR+SVM 基準。整體最佳表現出現在刑事附帶民事資料集之 no-leakage BoW + SVM，Macro F1 為 0.654。結果說明，在本研究任務中，經卡方篩選後的傳統稀疏文字表示仍具高度競爭力，且較複雜的 MNIR 特徵轉換未必能提升法律文本分類效果。"
    )
    add_para(doc, "關鍵字：智慧財產權、判決分類、文字特徵表示、BoW、TF-IDF、卡方特徵選擇、SVM")

    add_heading_numbered(doc, "第一章 緒論", level=1)
    add_heading_numbered(doc, "1.1 研究背景與動機", level=2)
    add_para(
        doc,
        "司法判決書包含大量事實認定、法律爭點與裁判結果資訊。隨著自然語言處理與機器學習方法逐漸應用於法律資料分析，如何將非結構化判決文字轉換為可比較、可重現且具解釋性的文字特徵，成為法律科技與實證法學研究的重要議題。智慧財產權案件兼具專業術語密集與案件類型多元特性，因此適合作為法律文本分類方法的實驗場域。"
    )
    add_heading_numbered(doc, "1.2 研究目的與問題", level=2)
    add_para(
        doc,
        "本研究不以單一複雜模型為核心，而是以特徵表示與分類流程比較為主軸。研究問題包括：BoW、TF 與 TF-IDF 何者較適合智慧財產判決分類；Direct SVM 是否優於 MNIR+SVM 基準流程；移除 outcome leakage 後模型表現是否改變；以及不同案件類型是否呈現不同分類難度。"
    )

    add_heading_numbered(doc, "第二章 文獻回顧", level=1)
    add_para(
        doc,
        "本章正式版本應補入法律文本分類、裁判預測、中文法律文本處理、文字特徵表示、卡方特徵選擇、SVM 與 MNIR 相關文獻。初稿建議將文獻回顧分為三個方向：第一，法律判決預測與實證法學中的文字探勘研究；第二，BoW、TF 與 TF-IDF 等傳統文字表示在分類任務中的角色；第三，特徵選擇與監督式特徵轉換方法對高維稀疏文本資料的影響。"
    )
    add_note(doc, "待補：正式文獻清單與引用格式。")

    add_heading_numbered(doc, "第三章 研究方法", level=1)
    add_heading_numbered(doc, "3.1 研究流程", level=2)
    add_para(
        doc,
        "本研究流程包含資料蒐集、案件與結果標籤標註、文本前處理、特徵表示、特徵選擇、模型訓練與模型評估。文本經 CKIP 斷詞後建立三種 DTM 表示：BoW、TF 與 TF-IDF。各表示分別進入 Direct SVM 與 MNIR+SVM 流程，以驗證特徵表示與特徵轉換方式對分類效果的影響。"
    )
    add_heading_numbered(doc, "3.2 資料集與標籤", level=2)
    add_para(
        doc,
        "有效案件依 JTYPE 區分為行政、民事、刑事與刑事附帶民事四類，裁判結果整理為 Win、Lose 與 Mixed。表 1 顯示各類案件之標籤分布，可見 Lose 類別在多數資料集中占比偏高，顯示本研究任務具明顯類別不平衡。"
    )
    add_table(doc, data["valid_summary"], "表 1 目標資料集之案件類型與裁判結果分布")

    add_heading_numbered(doc, "3.3 特徵表示與 Leakage 控制", level=2)
    add_para(
        doc,
        "本研究比較 BoW、TF 與 TF-IDF 三種特徵。BoW 保留詞彙出現次數，TF 對詞頻進行正規化，TF-IDF 則降低高頻常見詞影響。另為避免判決結果詞彙直接洩漏答案，本研究比較 with-leakage 與 no-leakage 兩種文本版本。"
    )
    add_heading_numbered(doc, "3.4 模型流程", level=2)
    add_para(
        doc,
        "Direct SVM 流程為：DTM → train-only 卡方特徵選擇 → SVM。MNIR+SVM 基準流程為：DTM → train-only 卡方特徵選擇 → MNIR 特徵轉換 → SVM。兩條流程均獨立以驗證集 Macro F1 選擇卡方特徵數 K 與 SVM 參數 C，因此 Direct SVM 不再沿用 MNIR 所選參數。"
    )
    add_heading_numbered(doc, "3.5 評估設計", level=2)
    add_para(
        doc,
        "資料採 70%、10%、20% 分層切分為訓練、驗證與測試集。模型調參只使用驗證集，測試集保留至最後評估。考量類別不平衡，本研究以 Macro F1 作為主要指標，並輔以 Accuracy、各類別 F1 與混淆矩陣。"
    )

    add_heading_numbered(doc, "第四章 實驗結果", level=1)
    add_heading_numbered(doc, "4.1 BoW、TF 與 TF-IDF 特徵比較", level=2)
    add_table(doc, data["feature_compare"], "表 2 Direct SVM 下三種特徵表示之最佳 Macro F1", percent_cols=["BoW", "TF", "TF-IDF"])
    add_para(
        doc,
        "結果顯示，最佳特徵表示會隨案件類型而異。行政與刑事附帶民事資料集以 BoW 表現較佳；民事資料集以 TF-IDF 表現較佳；刑事資料集則以 TF 表現略佳。此結果指出，單一特徵表示未必適用於所有法律案件類型。"
    )
    add_heading_numbered(doc, "4.2 Direct SVM 與 MNIR+SVM 比較", level=2)
    add_table(
        doc,
        data["direct_vs_mnir"],
        "表 3 Direct SVM 與 MNIR+SVM 於各資料集之最佳 Macro F1 比較",
        percent_cols=["Direct SVM Macro F1", "MNIR+SVM Macro F1", "差異"],
    )
    add_para(
        doc,
        "四類資料集中，Direct SVM 皆取得高於 MNIR+SVM 的最佳 Macro F1。此結果顯示，在本研究資料與標籤設定下，經卡方篩選後的原始稀疏詞彙特徵比 MNIR 轉換特徵更能保留與判決結果相關的分類訊息。"
    )
    add_heading_numbered(doc, "4.3 Outcome Leakage 影響", level=2)
    add_table(doc, data["leakage_compare"], "表 4 Direct SVM 下 with-leakage 與 no-leakage 最佳 Macro F1", percent_cols=["with leakage", "no leakage"])
    add_para(
        doc,
        "Leakage 移除對不同案件類型影響不一致。部分資料集在 no-leakage 版本仍可維持或取得較佳表現，表示模型並非完全依賴明顯裁判結果詞彙；但在部分設定中 with-leakage 表現較高，顯示 outcome leakage 仍可能影響模型評估。"
    )

    add_heading_numbered(doc, "4.4 參數高原與混淆矩陣", level=2)
    best_dir = data["best_run_dir"]
    add_figure(doc, best_dir / "direct_svm_chi2_k_svm_c_validation_heatmap.png", "圖 1 Direct SVM 參數表面範例：K × C 之 Validation Macro F1", width=5.8)
    add_figure(doc, best_dir / "model_comparison_test_metrics.png", "圖 2 最佳設定之模型比較圖", width=5.8)
    add_best_confusion_matrix(doc, data)

    add_heading_numbered(doc, "第五章 討論", level=1)
    add_heading_numbered(doc, "5.1 主要發現", level=2)
    add_para(
        doc,
        "本研究主要發現為：第一，Direct SVM 在四類資料集均優於 MNIR+SVM 基準；第二，BoW、TF 與 TF-IDF 的最佳表現依案件類型而異；第三，Accuracy 受多數類別影響明顯，Macro F1 更能揭示少數類別分類困難。"
    )
    add_heading_numbered(doc, "5.2 MNIR+SVM 作為基準之意義", level=2)
    add_para(
        doc,
        "MNIR+SVM 在本研究中作為基準流程，用以檢驗監督式特徵轉換是否能改善傳統文字特徵。實驗結果未支持此假設，可能原因包括法律文本中具有判別力的詞彙已可由稀疏特徵直接捕捉，MNIR 轉換反而壓縮或平滑了少數類別的重要訊息。"
    )
    add_heading_numbered(doc, "5.3 研究限制", level=2)
    add_para(
        doc,
        "本研究仍受標籤規則品質、類別不平衡、缺乏外部驗證資料與尚未納入大型語言模型等限制。後續可進一步檢查標籤正確性、加入 class weighting 或 resampling，並比較法律語言模型與傳統稀疏特徵模型之差異。"
    )

    add_heading_numbered(doc, "第六章 結論", level=1)
    add_para(
        doc,
        "本研究以智慧財產權判決書為對象，比較 BoW、TF、TF-IDF 三種特徵表示與 Direct SVM、MNIR+SVM 兩種流程。結果顯示，Direct SVM 在各資料集中均優於 MNIR+SVM，且最佳特徵表示依案件類型而異。此發現支持傳統稀疏文字表示在法律文本分類中仍具高度實用性，也提醒研究者在導入較複雜特徵轉換前，應建立嚴謹且獨立調參的 baseline。"
    )

    add_heading_numbered(doc, "參考文獻與待補事項", level=1)
    for item in [
        "補入法律文本分類與判決預測相關文獻。",
        "補入 BoW、TF、TF-IDF、卡方特徵選擇與 SVM 相關文獻。",
        "確認是否需要英文摘要、學校格式與正式章節編號。",
        "決定是否將所有 24 組模型結果放入附錄。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
